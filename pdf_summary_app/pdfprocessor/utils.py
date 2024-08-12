import PyPDF2
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from io import BytesIO


def extract_text_from_pdf(pdf_path):
    """Извлечение текста из текстового PDF."""
    text = ""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in range(len(reader.pages)):
            text += reader.pages[page].extract_text()
    return text


def extract_text_from_images(pdf_path):
    """Извлечение текста из PDF, содержащего изображения."""
    images = convert_from_path(pdf_path)
    text = ""
    for image in images:
        text += pytesseract.image_to_string(image)
    return text


def create_summary_docx(summary):
    """Создание документа .docx с текстом summary."""
    document = Document()
    document.add_heading('Ответ от Ollama', 0)
    document.add_paragraph(summary)

    # Сохраняем документ в поток
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream
